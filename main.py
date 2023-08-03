import json
from docx import Document



def progon_telegi(path_for_json, num):


    '''

    :param path_for_json: ссылка на файл
    :return: документ ворд

    '''

    doc = Document()

    with open(path_for_json) as data_table:
        data_for_table = json.load(data_table)

        table_name = data_for_table['name']
        table_load_type = data_for_table['load_type']

        doc.add_paragraph('Название: ')
        doc.add_paragraph('Тип нагрузки: ')

        table_vehicles = data_for_table['vehicles']

        telega = 0

        for data_vich in table_vehicles:

            telega += 1

            centr_posit_x_vinch = round(data_vich['x'], 2)
            centr_posit_y_vinch = round(data_vich['y'], 2)
            centr_posit_z_vinch = data_vich['z']

            table_of_pos_wheels = data_vich['wheels']


            creat_table(doc, telega,
                    centr_posit_x_vinch, centr_posit_y_vinch,
                    table_of_pos_wheels, table_load_type)


    doc.save(f'TMP_word/По порядку как загружали {num}.docx')

    return doc

def creat_table(doc, telega,
                centr_posit_x_vinch, centr_posit_y_vinch,
                table_of_pos_wheels, table_load_type):

    list_names_col = ['Ось/колесо', 'Х', 'Y', 'Z', 'Вес, тс', 'Воздействие от тележки нормативное']
    column = len(list_names_col)

    doc.add_paragraph(f'№ колонны: {telega}')
    doc.add_paragraph(f'Положение оси колонны относительно левого бокового ОБ, м {centr_posit_x_vinch}')
    doc.add_paragraph(f'Положение тележки вдоль проезда, м: {centr_posit_y_vinch}')

    table = doc.add_table(rows=1, cols=column)

    head_row = table.rows[0]
    for i, name in enumerate(list_names_col):
        head_row.cells[i].text = name

    for data_wheels in table_of_pos_wheels:

        wheel_position = data_wheels['position']
        wheel_x = round(data_wheels['x'], 2)
        wheel_y = round(data_wheels['y'], 2)
        wheel_z = round(data_wheels['z'], 2)

        print(table_load_type)
        ves_ts = 14
        result = ves_ts * wheel_z

        list_wheel = [wheel_position, wheel_x, wheel_y, wheel_z, ves_ts, result]

        cells = table.add_row().cells

        for i, data_list_wheel in enumerate(list_wheel):
            cells[i].text = str(data_list_wheel)